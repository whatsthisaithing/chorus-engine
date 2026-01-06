"""Document loader service for parsing various file formats.

Supports: PDF, CSV, Excel (XLSX), TXT, DOCX, Markdown
"""

import logging
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Any
from dataclasses import dataclass
import mimetypes

logger = logging.getLogger(__name__)


@dataclass
class LoadedDocument:
    """Container for loaded document content and metadata."""
    content: str
    metadata: Dict[str, Any]
    page_count: Optional[int] = None
    
    def __post_init__(self):
        """Ensure metadata dict exists."""
        if self.metadata is None:
            self.metadata = {}


class DocumentLoader:
    """Service for loading and parsing documents."""
    
    SUPPORTED_EXTENSIONS = {
        'pdf': 'application/pdf',
        'csv': 'text/csv',
        'xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        'xls': 'application/vnd.ms-excel',
        'txt': 'text/plain',
        'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'doc': 'application/msword',
        'md': 'text/markdown',
        'markdown': 'text/markdown',
    }
    
    def __init__(self):
        """Initialize document loader."""
        self._pdf_available = self._check_pdf_support()
        self._docx_available = self._check_docx_support()
        self._excel_available = self._check_excel_support()
        
        logger.info(f"DocumentLoader initialized (PDF: {self._pdf_available}, DOCX: {self._docx_available}, Excel: {self._excel_available})")
    
    def _check_pdf_support(self) -> bool:
        """Check if PyPDF2 is available."""
        try:
            import PyPDF2
            return True
        except ImportError:
            logger.warning("PyPDF2 not installed - PDF support disabled")
            return False
    
    def _check_docx_support(self) -> bool:
        """Check if python-docx is available."""
        try:
            import docx
            return True
        except ImportError:
            logger.warning("python-docx not installed - DOCX support disabled")
            return False
    
    def _check_excel_support(self) -> bool:
        """Check if openpyxl is available."""
        try:
            import openpyxl
            return True
        except ImportError:
            logger.warning("openpyxl not installed - Excel support disabled")
            return False
    
    def get_supported_extensions(self) -> List[str]:
        """
        Get list of currently supported file extensions.
        
        Returns:
            List of supported extensions (without dots)
        """
        supported = ['txt', 'md', 'markdown', 'csv']
        
        if self._pdf_available:
            supported.append('pdf')
        if self._docx_available:
            supported.extend(['docx', 'doc'])
        if self._excel_available:
            supported.extend(['xlsx', 'xls'])
        
        return supported
    
    def is_supported(self, file_path: str) -> bool:
        """
        Check if file type is supported.
        
        Args:
            file_path: Path to file
            
        Returns:
            True if supported, False otherwise
        """
        ext = Path(file_path).suffix.lstrip('.').lower()
        return ext in self.get_supported_extensions()
    
    def detect_file_type(self, file_path: str) -> Tuple[str, Optional[str]]:
        """
        Detect file type from extension and MIME type.
        
        Args:
            file_path: Path to file
            
        Returns:
            Tuple of (extension, mime_type)
        """
        path = Path(file_path)
        ext = path.suffix.lstrip('.').lower()
        
        # Get MIME type
        mime_type = mimetypes.guess_type(file_path)[0]
        if not mime_type and ext in self.SUPPORTED_EXTENSIONS:
            mime_type = self.SUPPORTED_EXTENSIONS[ext]
        
        return ext, mime_type
    
    def load(self, file_path: str) -> LoadedDocument:
        """
        Load document from file.
        
        Args:
            file_path: Path to document file
            
        Returns:
            LoadedDocument with content and metadata
            
        Raises:
            ValueError: If file type not supported
            FileNotFoundError: If file doesn't exist
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        ext, mime_type = self.detect_file_type(file_path)
        
        if not self.is_supported(file_path):
            raise ValueError(
                f"Unsupported file type: {ext}. "
                f"Supported types: {', '.join(self.get_supported_extensions())}"
            )
        
        logger.info(f"Loading document: {path.name} ({ext})")
        
        # Route to appropriate loader
        if ext == 'pdf':
            return self._load_pdf(path)
        elif ext in ('csv',):
            return self._load_csv(path)
        elif ext in ('xlsx', 'xls'):
            return self._load_excel(path)
        elif ext in ('txt', 'md', 'markdown'):
            return self._load_text(path)
        elif ext in ('docx', 'doc'):
            return self._load_docx(path)
        else:
            raise ValueError(f"No loader implemented for: {ext}")
    
    def _load_text(self, path: Path) -> LoadedDocument:
        """Load plain text or markdown file."""
        try:
            # Try UTF-8 first
            content = path.read_text(encoding='utf-8')
        except UnicodeDecodeError:
            # Fallback to latin-1
            logger.warning(f"UTF-8 decode failed for {path.name}, trying latin-1")
            content = path.read_text(encoding='latin-1')
        
        metadata = {
            'file_type': path.suffix.lstrip('.').lower(),
            'encoding': 'utf-8',
            'line_count': len(content.splitlines())
        }
        
        logger.info(f"Loaded text file: {len(content)} characters, {metadata['line_count']} lines")
        return LoadedDocument(content=content, metadata=metadata)
    
    def _load_csv(self, path: Path) -> LoadedDocument:
        """Load CSV file."""
        import csv
        
        rows = []
        with path.open('r', encoding='utf-8-sig', newline='') as f:
            reader = csv.DictReader(f)
            headers = reader.fieldnames
            
            for row in reader:
                rows.append(row)
        
        # Convert to readable text format
        content_parts = []
        
        # Add headers
        content_parts.append("CSV Data with columns: " + ", ".join(headers))
        content_parts.append("=" * 50)
        content_parts.append("")
        
        # Add rows (limit to reasonable number for context)
        for i, row in enumerate(rows[:1000], 1):
            row_text = f"Row {i}:"
            for key, value in row.items():
                row_text += f"\n  {key}: {value}"
            content_parts.append(row_text)
            content_parts.append("")
        
        if len(rows) > 1000:
            content_parts.append(f"... ({len(rows) - 1000} more rows omitted)")
        
        content = "\n".join(content_parts)
        
        metadata = {
            'file_type': 'csv',
            'row_count': len(rows),
            'column_count': len(headers),
            'columns': list(headers)
        }
        
        logger.info(f"Loaded CSV: {len(rows)} rows, {len(headers)} columns")
        return LoadedDocument(content=content, metadata=metadata)
    
    def _load_excel(self, path: Path) -> LoadedDocument:
        """Load Excel file."""
        if not self._excel_available:
            raise RuntimeError("openpyxl not installed - cannot load Excel files")
        
        import openpyxl
        
        workbook = openpyxl.load_workbook(path, data_only=True)
        sheet_names = workbook.sheetnames
        
        content_parts = []
        content_parts.append(f"Excel Workbook: {path.name}")
        content_parts.append(f"Sheets: {', '.join(sheet_names)}")
        content_parts.append("=" * 50)
        content_parts.append("")
        
        total_rows = 0
        
        for sheet_name in sheet_names:
            sheet = workbook[sheet_name]
            
            content_parts.append(f"Sheet: {sheet_name}")
            content_parts.append("-" * 30)
            
            # Get headers from first row
            headers = [cell.value for cell in sheet[1]]
            
            # Add rows (limit per sheet)
            rows_added = 0
            for row_idx, row in enumerate(sheet.iter_rows(min_row=2, values_only=True), 2):
                if rows_added >= 500:  # Limit per sheet
                    break
                
                row_text = f"Row {row_idx}:"
                for header, value in zip(headers, row):
                    if value is not None:
                        row_text += f"\n  {header}: {value}"
                
                content_parts.append(row_text)
                content_parts.append("")
                rows_added += 1
                total_rows += 1
            
            if sheet.max_row > 501:
                content_parts.append(f"... ({sheet.max_row - 501} more rows omitted)")
            
            content_parts.append("")
        
        content = "\n".join(content_parts)
        
        metadata = {
            'file_type': 'xlsx',
            'sheet_count': len(sheet_names),
            'sheet_names': sheet_names,
            'total_rows_processed': total_rows
        }
        
        logger.info(f"Loaded Excel: {len(sheet_names)} sheets, ~{total_rows} rows")
        return LoadedDocument(content=content, metadata=metadata)
    
    def _load_pdf(self, path: Path) -> LoadedDocument:
        """Load PDF file."""
        if not self._pdf_available:
            raise RuntimeError("PyPDF2 not installed - cannot load PDF files")
        
        import PyPDF2
        
        with path.open('rb') as f:
            reader = PyPDF2.PdfReader(f)
            page_count = len(reader.pages)
            
            # Extract text from all pages
            pages_text = []
            for page_num in range(page_count):
                page = reader.pages[page_num]
                text = page.extract_text()
                if text.strip():
                    pages_text.append(f"[Page {page_num + 1}]\n{text}")
            
            content = "\n\n".join(pages_text)
            
            # Extract metadata
            pdf_metadata = reader.metadata if reader.metadata else {}
            metadata = {
                'file_type': 'pdf',
                'page_count': page_count,
                'title': pdf_metadata.get('/Title', ''),
                'author': pdf_metadata.get('/Author', ''),
                'subject': pdf_metadata.get('/Subject', ''),
                'creator': pdf_metadata.get('/Creator', ''),
            }
            
            # Clean up empty metadata fields
            metadata = {k: v for k, v in metadata.items() if v}
        
        logger.info(f"Loaded PDF: {page_count} pages, {len(content)} characters")
        return LoadedDocument(content=content, metadata=metadata, page_count=page_count)
    
    def _load_docx(self, path: Path) -> LoadedDocument:
        """Load DOCX file."""
        if not self._docx_available:
            raise RuntimeError("python-docx not installed - cannot load DOCX files")
        
        import docx
        
        doc = docx.Document(path)
        
        # Extract paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        
        # Extract tables
        tables_text = []
        for table_idx, table in enumerate(doc.tables, 1):
            tables_text.append(f"[Table {table_idx}]")
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                tables_text.append(row_text)
            tables_text.append("")
        
        # Combine content
        content_parts = paragraphs
        if tables_text:
            content_parts.append("\n\n=== Tables ===\n")
            content_parts.extend(tables_text)
        
        content = "\n\n".join(content_parts)
        
        # Extract core properties
        props = doc.core_properties
        metadata = {
            'file_type': 'docx',
            'paragraph_count': len(paragraphs),
            'table_count': len(doc.tables),
            'title': props.title or '',
            'author': props.author or '',
            'subject': props.subject or '',
        }
        
        # Clean up empty metadata
        metadata = {k: v for k, v in metadata.items() if v}
        
        logger.info(f"Loaded DOCX: {len(paragraphs)} paragraphs, {len(doc.tables)} tables")
        return LoadedDocument(content=content, metadata=metadata)
